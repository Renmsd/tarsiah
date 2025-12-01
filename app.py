#app
# -*- coding: utf-8 -*-
import os
import json
from flask import Flask, render_template, request, send_file, jsonify, session
from docxtpl import DocxTemplate
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional
from workflow.rfp_workflow import build_rfp_graph
from evaluation_engine.evaluator import evaluate_proposal
from routes.compare_routes import compare_bp

import requests

RFP_GENERATOR_URL = "https://rfp-generator-production.up.railway.app/generate_rfp"  # عدّليها برابطك


# ============================================================
# ⚙️ Flask configuration
# ============================================================
app = Flask(__name__)
app.secret_key = "smart-rfp-ai-key"

from routes.table_routes import table_bp  # noqa: E402
app.register_blueprint(table_bp)
app.register_blueprint(compare_bp)



def fix_rtl_bullets(text: str) -> str:
    """Fixes the direction of bullets and punctuation for generated Arabic text."""
    if not isinstance(text, str):
        return text
    rtl_start = "\u202B"
    rtl_end = "\u202C"
    replacements = {
        "•": f"{rtl_start}•{rtl_end}",
        "-": f"{rtl_start}-{rtl_end}",
        "–": f"{rtl_start}–{rtl_end}",
        "—": f"{rtl_start}—{rtl_end}",
        ":": f"{rtl_start}:{rtl_end}",
        "؛": f"{rtl_start}؛{rtl_end}",
        ".": f"{rtl_start}.{rtl_end}",
        "،": f"{rtl_start}،{rtl_end}",
        "(": f"{rtl_start}({rtl_end}",
        ")": f"{rtl_start}){rtl_end}",
        "[": f"{rtl_start}[{rtl_end}",
        "]": f"{rtl_start}]{rtl_end}",
        "{": f"{rtl_start}{{{rtl_end}",
        "}": f"{rtl_start}}}{rtl_end}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


@app.route('/')
def home() -> str:
    return render_template('home.html')


@app.route('/rfp_input')
def rfp_input() -> str:
    return render_template('rfp_input.html')


@app.route('/rfp_generate', methods=['POST'])
def generate():
    import traceback

    print("\n==============================")
    print("🚀 /rfp_generate START")
    print("==============================")

    # 1) Collect user form data
    user_data = request.form.to_dict(flat=False)

    # حوّلي القوائم إلى نصوص مفصولة بفواصل
    for key, value in user_data.items():
        if isinstance(value, list):
            user_data[key] = "، ".join(value)

    print("📌 USER DATA:")
    print(user_data)

    # 2) include_sections (checkboxes)
    include_sections = {
        "Joint_Venture": request.form.get("include_Joint_Venture") is not None,
        "Tender_Split_Section": request.form.get("include_Tender_Split_Section") is not None,
        "Alternative_Offers": request.form.get("include_Alternative_Offers") is not None,
        "Insurance": request.form.get("include_Insurance") is not None,
    }

    print("📌 INCLUDE SECTIONS:")
    print(include_sections)

    # 3) Payload to Railway microservice
    payload = {
        "raw_input": user_data,
        "include_sections": include_sections,
    }

    print("\n📤 SENDING TO RAILWAY:")
    print("URL:", RFP_GENERATOR_URL)
    print("JSON:", payload)

    # 4) Call Railway ONCE
    try:
        resp = requests.post(
            RFP_GENERATOR_URL,
            json=payload,
            timeout=300,
        )

        print("\n📥 RAILWAY RESPONSE STATUS:", resp.status_code)
        print("📥 RAW RESPONSE TEXT:\n", resp.text)

        resp.raise_for_status()

        try:
            data = resp.json()
        except Exception:
            print("❌ Failed to parse JSON from Railway response")
            return render_template("rfp_generate.html", decisions={}, user_data=user_data)

        decisions = data.get("decisions", {})
        success = data.get("success", True)

        if not success:
            print("⚠️ Railway returned error:", data.get("error"))
            return render_template("rfp_generate.html", decisions={}, user_data=user_data)

    except Exception as e:
        print("\n❌ ERROR WHILE CALLING RAILWAY")
        traceback.print_exc()
        return render_template("rfp_generate.html", decisions={}, user_data=user_data)

    print("\n📌 FINAL DECISIONS FROM RAILWAY:")
    print(decisions)

    if not decisions:
        print("⚠️ Railway returned EMPTY decisions!")
        return render_template("rfp_generate.html", decisions={}, user_data=user_data)

    # 5) Build structure for template
    filtered_decisions = {
        k: {"value": v, "type": "llm"} for k, v in decisions.items()
    }

    # خزنّا في السيشن عشان /save
    session["user_data"] = user_data
    session["decisions"] = {k: v["value"] for k, v in filtered_decisions.items()}

    print("✅ FINISHED /rfp_generate")
    print("==============================\n")

    return render_template("rfp_generate.html", decisions=filtered_decisions, user_data=user_data)



@app.route('/save', methods=['POST'])
def save():
    session_user = session.get("user_data", {})
    session_llm = session.get("decisions", {})
    edited_data = request.form.to_dict()
    context = {**session_user, **session_llm, **edited_data}

    tpl = DocxTemplate("templates/rfp_general.docx")
    tpl.render(context)

    TABLE_KEYS = {
        "Bill_of_Quantities_and_Prices",
        "Materials_Specifications_Table",
        "Equipment_Specifications_Table",
        "Workers_Table",
    }
    safe_context = {k: v for k, v in context.items() if k not in TABLE_KEYS}

    llm_fields = {
        "Competition_Definition",
        "Text_of_Costs_of_Competition_Documents",
        "Project_Scope_of_Work",
        "Technical_Proposal_Documents",
        "Financial_Proposal_Documents",
        "Bid_Evaluation_Criteria",
        "Regulatory_Records_and_Licenses",
        "Inquiries_Section",
        "Tender_Split_Section",
        "Penalties",
        "Delay_Penalties",
        "Insurance",
        "Service_Delivery_Plan",
        # "Bill_of_Quantities_and_Prices",
        # "Materials_Specifications_Table",
        # "Equipment_Specifications_Table",
        "Service_Execution_Method",
        "Alternative_Offers",
        "Offer_Formatting_Requirements",
    }

    for key in list(safe_context.keys()):
        if key in llm_fields and isinstance(safe_context[key], str):
            safe_context[key] = fix_rtl_bullets(safe_context[key])

    tpl.render(safe_context)
    output_folder = os.path.join(os.path.expanduser("~"), "Documents", "RFP_outputs")
    os.makedirs(output_folder, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_folder, f"filled_{timestamp}.docx")
    tpl.save(output_path)

    def insert_table_if_exists(session_key: str,
                               placeholder_text: str,
                               heading_text: Optional[str] = None) -> None:
        """إدراج جدول داخل المستند في موقع محدد.

        إذا تم تمرير heading_text، سيُدرج الجدول مباشرة بعد فقرة تحتوي على العنوان.
        إذا لم يوجد العنوان، سيبحث عن placeholder ويُدرج بعده. وإذا لم يُعثر على أي منهما، يتخطى الإدراج.
        """
        if session_key not in session:
            return
        table_text = session[session_key]
        lines = [l.strip() for l in table_text.split("\n") if "|" in l]
        if lines:
            headers = [h.strip() for h in lines[0].split("|")]
            rows = [l.split("|") for l in lines[1:]]
        if not lines:
            return
        
        headers = [h.strip() for h in lines[0].split("|")]
        rows = [l.split("|") for l in lines[1:]]

        doc = Document(output_path)

        def find_paragraph_with_text(doc_: Document, text: str):
            for p in doc_.paragraphs:
                if text in p.text:
                    return p
            for table in doc_.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if text in p.text:
                                return p
            return None

        # البحث عن العنوان أولاً إذا تم تمريره
        target_paragraph = None
        if heading_text:
            for p in doc.paragraphs:
                if p.text.strip() == heading_text.strip():
                    target_paragraph = p
                    break
            if not target_paragraph:
                target_paragraph = find_paragraph_with_text(doc, heading_text)
        if not target_paragraph:
            # البحث عن placeholder
            target_paragraph = find_paragraph_with_text(doc, placeholder_text)
        if not target_paragraph:
            print(f"⚠️ لم يتم العثور على العنوان '{heading_text}' ولا على placeholder {placeholder_text} داخل القالب.")
            return

        # إذا وجد placeholder داخل الفقرة، امسح نصه
        if placeholder_text and placeholder_text in target_paragraph.text:
            target_paragraph.text = ""

        # إنشاء الجدول وضبط اتجاهه RTL
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        tbl = table._element
        tbl.set(qn("w:tblDir"), "rtl")
        tbl.set(qn("w:tblLayout"), "fixed")

        # تنسيق صف العناوين
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9D9D9")
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

        # صفوف البيانات
        for r in rows:
            row_cells = table.add_row().cells
            for i, c in enumerate(r):
                p = row_cells[i].paragraphs[0]
                run = p.add_run(c.strip())
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # إدراج الجدول مباشرة بعد الفقرة الهدف
        target_paragraph._element.addnext(table._element)
        doc.save(output_path)
        print(f"✅ تم إدراج الجدول: {session_key}")

    insert_table_if_exists(
        "Bill_of_Quantities_and_Prices",
        "{{Bill_of_Quantities_and_Prices}}",
        heading_text="جدول الكميات والأسعار",
    )
    insert_table_if_exists(
        "Materials_Specifications_Table",
        "{{Materials_Specifications_Table}}",
        heading_text="جدول مواصفات المواد",
    )
    insert_table_if_exists(
        "Equipment_Specifications_Table",
        "{{Equipment_Specifications_Table}}",
        heading_text="واصفات المعدات",
    )

    insert_table_if_exists(
        "Workers_Table",
        "{{Workers_Table}}",
        heading_text="ثانياً: جدول مواصفات فريق العمل",
    )


    project_name = context.get("Competition_Name") or "مشروع بدون اسم"
    current_date = datetime.now().strftime("%Y-%m-%d")

    json_path = os.path.join("static", "projects.json")

    # قراءة الملف أو إنشاؤه
    if os.path.exists(json_path):
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {"ar": [], "en": []}

    # حفظ بيانات المشروع + ملف الناتج
    data["ar"].insert(0, {
        "name": project_name,
        "date": current_date,
        "file": output_path   # ← مسار ملف الـ RFP
    })

    data["en"].insert(0, {
        "name": project_name,
        "date": current_date,
        "file": output_path
    })

# كتابة الملف من جديد
    os.makedirs(os.path.dirname(json_path), exist_ok=True)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


    session["generated_file"] = output_path
    return render_template("result.html", project_name=project_name, download_url="/download")


@app.route('/download')
def download_file():
    file_path = session.get("generated_file")
    if not file_path or not os.path.exists(file_path):
        return "⚠️ الملف غير موجود. يرجى إعادة التوليد.", 404
    return send_file(file_path, as_attachment=True)

@app.route("/comparison")
def comparison_page():
    return render_template("comparison.html")


if __name__ == "__main__":
    app.run(debug=True) 